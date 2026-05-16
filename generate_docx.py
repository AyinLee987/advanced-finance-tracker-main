"""
Generate a formatted Word document (report.docx) from report.md content.
Typography requirements:
  - Body: 12pt Calibri
  - H1 (Section Titles): 14pt Bold
  - H2 (Sub-sections): 12pt Bold
  - Code: 10pt Consolas
Citation format: [1], [2]
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Global style defaults ──────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Helper: set font for a run ─────────────────────────────────────────────
def set_run_font(run, name='Calibri', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)

def add_heading_custom(text, level=1):
    """Add a heading with custom formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, 'Calibri', 14, bold=True)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, 'Calibri', 12, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p

def add_body(text):
    """Add a normal body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 'Calibri', 12)
    return p

def add_code_block(code_text):
    """Add a code block with monospaced font."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    set_run_font(run, 'Consolas', 10)
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F2F2F2'
    })
    shading.append(shd)
    return p

def add_table_row(table, cells_data, bold=False):
    """Add a row to a table."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_run_font(run, 'Calibri', 11, bold=bold)
    return row

# ═══════════════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('XSS Vulnerability Fix Report:\nHTML Escaping in Advanced Finance Tracker')
set_run_font(title_run, 'Calibri', 16, bold=True)
title_p.paragraph_format.space_after = Pt(12)

# ═══════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════
add_heading_custom('1. Introduction', level=1)
add_body(
    'Cross-Site Scripting (XSS) remains one of the most prevalent web security vulnerabilities, '
    'ranking second in the CWE Top 25 list for 2022 [1]. This report documents the '
    'XSS vulnerability fix implemented in the Advanced Finance Tracker project, a personal finance '
    'management application built with vanilla JavaScript. The fix introduces an escapeHtml function '
    'that sanitizes user-generated content before rendering it in the DOM, effectively preventing '
    'stored XSS attacks.'
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2: DEFICIENCY ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_custom('Section 2: Deficiency Analysis', level=1)

# --- 2.1 Detection ---
add_heading_custom('2.1 Detection', level=2)
add_body(
    'The XSS vulnerability was identified through manual code review. The renderTransactionItem '
    'function in main.js (line 467) constructs HTML via template literals that interpolate '
    'user-supplied data \u2014 tx.title and tx.category \u2014 directly into the DOM through innerHTML. '
    'An attacker can inject <img src=x onerror="alert(\'XSS\')"> as a transaction title. When rendered, '
    'the browser parses the <img> tag and executes the onerror handler \u2014 a classic Stored XSS attack. '
    'No sanitization was present at any rendering point.'
)

# --- 2.2 Literature ---
add_heading_custom('2.2 Literature', level=2)
add_body(
    'Two sources guided the remediation. Tang et al. [1] identify XSS root cause as the browser\u2019s '
    'inability to distinguish legitimate code from injected code. Their paper demonstrates (Section III-B, '
    'Fig. 4\u20135) that replacing innerHTML with textContent \u2014 or equivalently, escaping HTML entities '
    'before DOM insertion \u2014 prevents the browser from parsing special symbols in tags, neutralizing '
    'malicious code. This is illustrated through MISP vulnerability CVE-2023-24027, where the fix replaced '
    'innerHTML with textContent on line 165 of action_table.js. Cloudflare\u2019s XSS prevention documentation [2] '
    'reinforces this, emphasizing that character escaping \u2014 converting <, >, &, ", and \' to HTML entities '
    '\u2014 is a foundational defense. Both sources agree: treat all user data as untrusted text, not executable code.'
)

# --- 2.3 Implementation ---
add_heading_custom('2.3 Implementation', level=2)
add_body('Guided by the literature, the fix introduces an escapeHtml function that sanitizes user data at render time:')

add_code_block(
    'const escapeHtml = (str) =>\n'
    '  String(str)\n'
    '    .replace(/&/g, "&")\n'
    '    .replace(/</g, "<")\n'
    '    .replace(/>/g, ">")\n'
    '    .replace(/"/g, """)\n'
    "    .replace(/'/g, \"&#039;\");"
)

add_body(
    'The function is applied at every DOM insertion point. In renderTransactionItem, the title and '
    'category are wrapped: const safeTitle = escapeHtml(tx.title). Month-group labels in renderTransactions '
    'are similarly sanitized. The fix exists in both main.js (lines 8\u201314) and utils.js (lines 4\u201310).'
)

# Before vs After table
p_table_title = doc.add_paragraph()
run_table_title = p_table_title.add_run('Before vs. After code comparison:')
set_run_font(run_table_title, 'Calibri', 12, bold=True)
p_table_title.paragraph_format.space_before = Pt(8)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'

header = table.rows[0]
for i, text in enumerate(['Aspect', 'Before (Vulnerable)', 'After (Protected)']):
    cell = header.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, 'Calibri', 11, bold=True)

add_table_row(table, ['Rendering', '<p>${tx.title}</p>', '<p>${escapeHtml(tx.title)}</p>'])
add_table_row(table, ['Input', '<img src=x onerror="alert(1)">', '<img src=x onerror="alert(1)">'])
add_table_row(table, ['Output', 'Image rendered, script executes', '<img src=x onerror="alert(1)"> as text'])

add_body(
    'The fix is tested in main.test.js (lines 27\u201361) with 8 test cases covering < > escaping, '
    'ampersand, quotes, XSS payloads, null/number coercion, and safe strings. This aligns with '
    'Tang et al.\u2019s defense-in-depth principle [1] and Cloudflare\u2019s best practices [2].'
)

# ═══════════════════════════════════════════════════════════════════════════
# 3. REFERENCES
# ═══════════════════════════════════════════════════════════════════════════
add_heading_custom('3. References', level=1)

ref1 = (
    '[1] Tang, C., Wang, Q., Cheng, G., Liang, H., Peng, J., Yang, M., Liu, W., Liu, M., & Sha, L. (2024). '
    'A Cross-Site Scripting Attack Protection Framework Based on Managed Proxy. '
    '2024 IEEE 23rd International Conference on Trust, Security and Privacy in Computing and Communications '
    '(TrustCom), pp. 1896\u20131903. DOI: 10.1109/TrustCom63139.2024.00262.'
)
add_body(ref1)

ref2 = (
    '[2] Cloudflare. How to prevent XSS attacks. '
    'Available at: https://www.cloudflare.com/learning/security/how-to-prevent-xss-attacks/'
)
add_body(ref2)

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX: SUPPORTING FIGURES
# ═══════════════════════════════════════════════════════════════════════════
add_heading_custom('Appendix: Supporting Figures', level=1)

# Figure 1
fig1_caption = doc.add_paragraph()
run_fig1 = fig1_caption.add_run('Figure 1: MISP Vulnerability Disclosure (CVE-2023-24027) \u2014 Tang et al. [1] Fig. 4')
set_run_font(run_fig1, 'Calibri', 11, bold=True)

fig1_path = os.path.join(os.path.dirname(__file__), 'assets', 'fig4_misp_vulnerability.png')
if os.path.exists(fig1_path):
    doc.add_picture(fig1_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

fig1_desc = doc.add_paragraph()
run_fig1d = fig1_desc.add_run(
    'This figure from Tang et al. [1] shows the MISP vulnerability CVE-2023-24027, where innerHTML '
    'renders user-controlled data (network address names), creating an XSS vector. This mirrors the '
    'vulnerability found in our application before the fix.'
)
set_run_font(run_fig1d, 'Calibri', 11, italic=True)

# Figure 2
fig2_caption = doc.add_paragraph()
run_fig2 = fig2_caption.add_run('Figure 2: MISP Vulnerability Fix \u2014 Tang et al. [1] Fig. 5')
set_run_font(run_fig2, 'Calibri', 11, bold=True)

fig2_path = os.path.join(os.path.dirname(__file__), 'assets', 'fig5_misp_fix.png')
if os.path.exists(fig2_path):
    doc.add_picture(fig2_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

fig2_desc = doc.add_paragraph()
run_fig2d = fig2_desc.add_run(
    'This figure illustrates the fix: replacing innerHTML with textContent on line 165 of action_table.js, '
    'preventing the browser from parsing special symbols in tags. Our escapeHtml function achieves the '
    'same protective effect by encoding HTML entities before DOM insertion.'
)
set_run_font(run_fig2d, 'Calibri', 11, italic=True)

# ── Save ───────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), 'report.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
